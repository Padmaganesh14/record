from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import base64
import io
import uuid
import os
import tempfile

app = Flask(__name__)
CORS(app)

app.config['MAX_CONTENT_LENGTH'] = 200 * 1024 * 1024
STORAGE_FILE = 'templates.json'
TEMPLATE_STORAGE = {}


def load_templates():
    """Load templates from storage file on startup."""
    global TEMPLATE_STORAGE
    if os.path.exists(STORAGE_FILE):
        try:
            with open(STORAGE_FILE, 'r') as f:
                TEMPLATE_STORAGE = json.load(f)
        except Exception:
            TEMPLATE_STORAGE = {}


def save_templates_to_file():
    """Save templates to storage file."""
    with open(STORAGE_FILE, 'w') as f:
        json.dump(TEMPLATE_STORAGE, f)


def add_page_border(section):
    """Add a border to the page."""
    xml = parse_xml(
        r'<w:pgBorders %s w:offsetFrom="page">'
        r'<w:top w:val="single" w:sz="12" w:space="24"/>'
        r'<w:left w:val="single" w:sz="12" w:space="24"/>'
        r'<w:bottom w:val="single" w:sz="12" w:space="24"/>'
        r'<w:right w:val="single" w:sz="12" w:space="24"/>'
        r'</w:pgBorders>' % nsdecls('w')
    )
    section._sectPr.append(xml)


def add_watermark(section, text):
    """Add a watermark to the document header."""
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    xml = f"""
    <w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <w:pPr><w:pStyle w:val="Header"/></w:pPr>
        <w:r>
            <w:rPr><w:noProof/></w:rPr>
            <w:pict>
                <v:shape xmlns:v="urn:schemas-microsoft-com:vml" id="WordWatermark" o:spaste="f" 
                         type="#_x0000_t75" style="position:absolute;margin-left:0;margin-top:0;width:500pt;height:200pt;z-index:-251658240;mso-wrap-edited:f;mso-position-horizontal:center;mso-position-horizontal-relative:margin;mso-position-vertical:center;mso-position-vertical-relative:margin">
                    <v:textpath on="t" string="{text}" style="font-family:'Calibri';font-size:36pt;color:#d3d3d3;opacity:0.5"/>
                </v:shape>
            </w:pict>
        </w:r>
    </w:p>
    """
    p._p.append(parse_xml(xml))


# Load templates on startup
load_templates()


# ============================================================================
# ROUTES
# ============================================================================

@app.route("/", methods=["GET"])
def health_check():
    """Health check endpoint."""
    return jsonify({
        "status": "ok",
        "message": "Word Record Generator API is running"
    }), 200


@app.route("/generate", methods=["POST"])
def generate():
    """Generate a Word document from blocks and settings."""
    try:
        doc = Document()
        section = doc.sections[0]

        # Page layout configuration
        layout = request.form.get("layout", "normal")
        if layout == "legal":
            section.page_height = Inches(14)
            section.page_width = Inches(8.5)
        elif layout == "a4":
            section.page_height = Inches(11.69)
            section.page_width = Inches(8.27)

        # Page borders
        if request.form.get("border") == "on":
            add_page_border(section)

        # Watermark
        watermark_text = request.form.get("watermark", "").strip()
        if watermark_text:
            add_watermark(section, watermark_text)

        # Process document blocks
        headings_json = request.form.get("headings", "[]")
        headings = json.loads(headings_json)

        for item in headings:
            if item.get("type") == "image" and item.get("content"):
                try:
                    img_data = base64.b64decode(item["content"].split(",")[1])
                    img_stream = io.BytesIO(img_data)
                    doc.add_picture(img_stream, width=Inches(4))
                    last_p = doc.paragraphs[-1]
                    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass
            else:
                p = doc.add_paragraph()
                run = p.add_run(item.get("text", ""))
                run.font.size = Pt(int(item.get("size", 12)))
                run.font.bold = item.get("bold", False)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if item.get(
                    "type") == "heading" else WD_ALIGN_PARAGRAPH.LEFT

        # Save to /tmp for cloud deployment compatibility
        filename = f"record_{uuid.uuid4().hex[:8]}.docx"
        filepath = os.path.join(tempfile.gettempdir(), filename)
        doc.save(filepath)

        # Return file with proper headers for mobile compatibility
        return send_file(
            filepath,
            as_attachment=True,
            download_name="record.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        print(f"Error generating document: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/save_template", methods=["POST"])
def save_template():
    """Save a document template."""
    try:
        data = request.json
        code = f"REC-{uuid.uuid4().hex[:6].upper()}"
        TEMPLATE_STORAGE[code] = data
        save_templates_to_file()
        return jsonify({"success": True, "code": code}), 200
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 400


@app.route("/load_template/<code>", methods=["GET"])
def load_template_route(code):
    """Load a saved template by code."""
    template = TEMPLATE_STORAGE.get(code.upper())
    if template:
        return jsonify({"success": True, "template": template}), 200
    return jsonify({"success": False, "message": "Template not found"}), 404


@app.route("/list_templates", methods=["GET"])
def list_templates():
    """List all saved templates."""
    templates_list = []
    for code, data in TEMPLATE_STORAGE.items():
        templates_list.append({
            "code": code,
            "heading_count": data.get("heading_count", 0),
            "created_at": "Recent"
        })
    return jsonify({"success": True, "templates": templates_list}), 200


@app.route("/delete_template/<code>", methods=["DELETE"])
def delete_template(code):
    """Delete a saved template."""
    if code.upper() in TEMPLATE_STORAGE:
        del TEMPLATE_STORAGE[code.upper()]
        save_templates_to_file()
        return jsonify({"success": True}), 200
    return jsonify({"success": False, "message": "Not found"}), 404


if __name__ == "__main__":
    # Bind to 0.0.0.0 to allow mobile devices to access via local IP
    # For development: http://YOUR_LOCAL_IP:5000
    # For production: Render will override these settings
    app.run(host="0.0.0.0", port=5000, debug=False)

